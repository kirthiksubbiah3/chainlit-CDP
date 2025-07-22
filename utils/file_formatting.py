from docx import Document as DocxDocument


def format_docx_content(content: str, output_path: str):
    lines = content.splitlines()
    doc = DocxDocument()
    for line in lines:
        if not line.strip():
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(output_path)
