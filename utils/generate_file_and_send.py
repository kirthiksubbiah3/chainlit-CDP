"""File generation and handling utilities"""

import os
import tempfile
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch

import chainlit as cl


async def generate_file_and_send(content: str, file_format: str, for_id: str):
    """Main function to generate and send file"""
    if not content.strip():
        return
    lines = content.splitlines()
    if file_format == "pdf":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            pdf_path = tmp.name
        try:
            doc = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72,
            )

            styles = getSampleStyleSheet()
            styles.add(
                ParagraphStyle(
                    name="Heading",
                    fontSize=14,
                    leading=16,
                    spaceAfter=12,
                    spaceBefore=12,
                    alignment=TA_LEFT,
                    fontName="Helvetica-Bold",
                )
            )
            styles.add(
                ParagraphStyle(
                    name="Body",
                    fontSize=11,
                    leading=14,
                    spaceAfter=8,
                    alignment=TA_LEFT,
                )
            )
            styles.add(
                ParagraphStyle(
                    name="MyBullet",
                    fontSize=11,
                    leading=14,
                    spaceAfter=6,
                    leftIndent=12,
                    bulletIndent=0,
                )
            )
            story = []
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    story.append(Spacer(1, 0.2 * inch))
                    continue
                if line.endswith(":"):
                    story.append(Paragraph(line, styles["Heading"]))
                elif line.startswith("- "):
                    story.append(Paragraph(line[2:].strip(), styles["MyBullet"]))
                else:
                    story.append(Paragraph(line, styles["Body"]))

            doc.build(story)

            await cl.Message(content="📄 Here is your styled PDF report:").send()
            await cl.File(name="response.pdf", path=pdf_path, display="inline").send(
                for_id=for_id
            )
        finally:
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

    elif file_format == "docx":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            docx_path = tmp.name
        try:
            doc = Document()
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph("")
                    continue
                doc.add_paragraph(line)

            doc.save(docx_path)

            await cl.Message(content="📝 Here is your DOCX report:").send()
            await cl.File(name="response.docx", path=docx_path, display="inline").send(
                for_id=for_id
            )
        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)
