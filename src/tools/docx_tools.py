"""LangChain tool to generate DOCX files from markdown content."""

import tempfile
from docx import Document
from bs4 import BeautifulSoup
import markdown2
from langchain.tools import tool
import chainlit as cl


@tool
def generate_docx(content: str, filename: str) -> str:
    """
    Use this tool when the user wants the report in a Word document or DOCX file format.
    or want to save the output as a DOC file.
    the filename should be relevant to the content.
    """
    # Convert markdown to HTML
    html = markdown2.markdown(content, extras=["fenced-code-blocks", "tables"])
    soup = BeautifulSoup(html, "html.parser")

    # Create DOCX document
    doc = Document()

    for element in soup.find_all(True):
        if element.name == "h1":
            doc.add_heading(element.text, level=1)
        elif element.name == "h2":
            doc.add_heading(element.text, level=2)
        elif element.name == "h3":
            doc.add_heading(element.text, level=3)
        elif element.name == "p":
            doc.add_paragraph(element.text)
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.text, style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.text, style="List Number")
        elif element.name == "pre":
            code = element.text
            doc.add_paragraph(code, style="Code")

    # Save temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)

    cl.user_session.set("file_path", temp_file.name)
    cl.user_session.set("file_name", f"{filename}.docx")
    return temp_file.name
